from pathlib import Path
from typing import Optional
import mimetypes
from fastapi import UploadFile
import uuid
import logging

logger = logging.getLogger(__name__)


def save_upload(file: UploadFile, upload_dir: Path) -> Path:
    """Persist an uploaded file to disk."""
    upload_dir.mkdir(parents=True, exist_ok=True)
    # Use a unique filename to avoid conflicts
    file_ext = Path(file.filename).suffix
    unique_filename = f"{uuid.uuid4()}{file_ext}"
    destination = upload_dir / unique_filename
    with destination.open("wb") as buffer:
        while chunk := file.file.read(1024 * 1024):
            buffer.write(chunk)
    return destination


def extract_text(path: Path) -> str:
    """
    Extract text from various file formats.
    Supports: PDF, DOCX, TXT, MD, RTF, and code files.
    """
    mime_type, _ = mimetypes.guess_type(str(path))
    suffix = path.suffix.lower()
    
    # Text files (.txt, .md, .markdown, etc.)
    if mime_type and mime_type.startswith('text/'):
        return path.read_text(encoding='utf-8', errors='ignore')
    
    # Markdown files
    if suffix in ['.md', '.markdown']:
        return path.read_text(encoding='utf-8', errors='ignore')
    
    # PDF files - try pdfplumber first (better extraction), fallback to PyPDF2
    if mime_type == 'application/pdf' or suffix == '.pdf':
        # Try pdfplumber first (better text extraction)
        try:
            import pdfplumber
            text_parts = []
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            if text_parts:
                return '\n\n'.join(text_parts)
        except ImportError:
            logger.debug("pdfplumber not available, trying PyPDF2")
        except Exception as e:
            logger.warning(f"pdfplumber failed for {path.name}: {e}, trying PyPDF2")
        
        # Fallback to PyPDF2
        try:
            import PyPDF2
            with path.open('rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                text_parts = []
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                if text_parts:
                    return '\n\n'.join(text_parts)
                return f"[PDF file {path.name} - No text content found]"
        except ImportError:
            return f"[PDF file {path.name} - PDF libraries not installed. Install with: pip install pdfplumber PyPDF2]"
        except Exception as e:
            return f"[Error reading PDF {path.name}: {str(e)}]"
    
    # DOCX files
    if (mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or 
        suffix == '.docx'):
        try:
            from docx import Document
            doc = Document(path)
            text_parts = []
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            return '\n\n'.join(text_parts) if text_parts else f"[DOCX file {path.name} - No text content found]"
        except ImportError:
            return f"[DOCX file {path.name} - python-docx not installed. Install with: pip install python-docx]"
        except Exception as e:
            return f"[Error reading DOCX {path.name}: {str(e)}]"
    
    # RTF files
    if mime_type == 'application/rtf' or suffix == '.rtf':
        try:
            import striprtf
            rtf_text = path.read_text(encoding='utf-8', errors='ignore')
            return striprtf.RTF(rtf_text).plain_text
        except ImportError:
            logger.warning("striprtf not available for RTF files")
            # Try to read as plain text (may have RTF markup)
            return path.read_text(encoding='utf-8', errors='ignore')
        except Exception as e:
            logger.warning(f"Error reading RTF {path.name}: {e}")
            return path.read_text(encoding='utf-8', errors='ignore')
    
    # CSV files - format as readable table
    if suffix == '.csv' or mime_type == 'text/csv':
        try:
            import csv
            text_parts = []
            with path.open('r', encoding='utf-8', errors='ignore', newline='') as f:
                reader = csv.reader(f)
                rows = list(reader)
                
                if not rows:
                    return f"[CSV file {path.name} - Empty file]"
                
                # Get headers (first row)
                headers = rows[0] if rows else []
                text_parts.append(f"CSV Data from {path.name}")
                text_parts.append(f"Columns: {', '.join(headers)}")
                text_parts.append(f"Total rows: {len(rows) - 1}")
                text_parts.append("")
                
                # Format as readable table (header row + data)
                text_parts.append(" | ".join(headers))
                text_parts.append("-" * 50)
                
                for row in rows[1:]:  # Skip header
                    # Pad row if needed
                    while len(row) < len(headers):
                        row.append("")
                    text_parts.append(" | ".join(row[:len(headers)]))
                
                return '\n'.join(text_parts)
        except Exception as e:
            logger.warning(f"CSV parsing failed for {path.name}: {e}, reading as text")
            return path.read_text(encoding='utf-8', errors='ignore')
    
    # Code files - read as plain text with syntax preserved
    code_extensions = {
        '.py', '.js', '.jsx', '.ts', '.tsx', '.java', '.cpp', '.c', '.h', '.hpp',
        '.cs', '.go', '.rs', '.rb', '.php', '.swift', '.kt', '.scala', '.r',
        '.sql', '.html', '.css', '.scss', '.sass', '.less', '.xml', '.json',
        '.yaml', '.yml', '.toml', '.ini', '.cfg', '.conf', '.sh', '.bash',
        '.zsh', '.ps1', '.bat', '.cmd', '.vue', '.svelte', '.dart', '.lua'
    }
    if suffix in code_extensions:
        return path.read_text(encoding='utf-8', errors='ignore')
    
    # DOC files (older format - requires additional library)
    if (mime_type == 'application/msword' or suffix == '.doc'):
        return f"[DOC file {path.name} - .doc format requires additional library. Please convert to .docx or .txt]"
    
    # Default: try to read as text
    try:
        return path.read_text(encoding='utf-8', errors='ignore')
    except Exception:
        return f"[Could not extract text from {path.name}]"


def read_file(path: Path, encoding: str = "utf-8") -> str:
    """Read a text file with fallback error handling."""
    return path.read_text(encoding=encoding, errors="ignore")
